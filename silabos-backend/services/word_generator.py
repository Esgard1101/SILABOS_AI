"""Generación DOCX y PDF para sílabos."""

import logging
import os
import re
from datetime import date, datetime, timedelta
from html import escape
from io import BytesIO
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def _val(v: Any, fallback: str = "") -> str:
    if v is None:
        return fallback
    s = str(v).strip()
    if not s or s in {"-", "-", "None", "none", "null", "NULL", "/"}:
        return fallback
    return s


def _pair(left: Any, right: Any, fallback: str = "-") -> str:
    parts = [_val(left), _val(right)]
    visible = [part for part in parts if part]
    return " / ".join(visible) if visible else fallback


def _prerequisite_value(value: Any) -> str:
    text = _val(value)
    if not text or re.fullmatch(r"(?i)(no\s+registrado|ninguno|sin\s+prerrequisito.*)", text):
        return "No aplica"
    return text


SEMESTER_WEEKS = 16


def _parse_iso_date(value: Any) -> date | None:
    text = _val(value)
    if not text:
        return None
    try:
        return datetime.strptime(text[:10], "%Y-%m-%d").date()
    except (TypeError, ValueError):
        return None


def _compute_end_date(start_value: Any, total_weeks: int = SEMESTER_WEEKS) -> str:
    start = _parse_iso_date(start_value)
    if not start:
        return ""
    return (start + timedelta(weeks=total_weeks)).strftime("%Y-%m-%d")


def _compute_week_dates(
    semester: Any,
    total_weeks: int = SEMESTER_WEEKS,
    start_date: Any = None,
) -> list[str]:
    explicit_start = _parse_iso_date(start_date)
    if explicit_start:
        return [(explicit_start + timedelta(weeks=i)).strftime("%Y-%m-%d") for i in range(total_weeks)]

    semester_text = _val(semester).upper()
    match = re.match(r"\s*(\d{4})\s*[-_]?\s*(I{1,2})\s*$", semester_text)
    if not match:
        return ["---"] * total_weeks

    year = int(match.group(1))
    period = match.group(2)
    target_month = 3 if period == "I" else 8
    first = date(year, target_month, 1)
    days_until_monday = (7 - first.weekday()) % 7
    first_monday = first + timedelta(days=days_until_monday)
    start = first_monday + timedelta(weeks=3)
    return [(start + timedelta(weeks=i)).strftime("%Y-%m-%d") for i in range(total_weeks)]


def _week_date_from_row(row: dict[str, Any], week_number: int, week_dates: list[str]) -> str:
    direct = _val(row.get("fecha") or row.get("date_range") or row.get("date"))
    if direct:
        return direct
    index = int(week_number) - 1
    if 0 <= index < len(week_dates):
        return _val(week_dates[index], "---")
    return "---"


def _format_hour_value(value: Any) -> str:
    formatted = _val(value)
    if not formatted:
        return ""

    normalized_text = formatted.replace(",", ".")
    if not re.fullmatch(r"\d+(\.\d+)?", normalized_text):
        return formatted

    numeric_value = float(normalized_text)
    if numeric_value < SEMESTER_WEEKS:
        return formatted

    weekly_value = numeric_value / SEMESTER_WEEKS
    if weekly_value.is_integer():
        return str(int(weekly_value))
    return f"{weekly_value:.2f}".rstrip("0").rstrip(".")


def _format_hours(theory: Any, practice: Any, fallback: str = "-") -> str:
    t = _format_hour_value(theory)
    p = _format_hour_value(practice)
    if t and p:
        return f"Teoría: {t} / Práctica: {p}"
    if t:
        return f"Teoría: {t}"
    if p:
        return f"Práctica: {p}"
    return fallback


def _sanitize_sentence_spacing(text: Any) -> str:
    raw = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", raw) if part.strip()]
    if not paragraphs:
        paragraphs = [raw.strip()] if raw.strip() else []
    cleaned: list[str] = []
    for paragraph in paragraphs:
        value = re.sub(r"\s+", " ", paragraph).strip()
        value = re.sub(r"(?<=[a-záéíóúñ])\.([A-ZÁÉÍÓÚÑ])", r". \1", value)
        value = re.sub(r"(?<=[a-záéíóúñ])([!?])([A-ZÁÉÍÓÚÑ])", r"\1 \2", value)
        if value:
            cleaned.append(value)
    return "\n\n".join(cleaned)


def _safe_filename(nombre_curso: str) -> str:
    nombre = re.sub(r"[^\w\s-]", "", nombre_curso or "silabo")
    return re.sub(r"\s+", "_", nombre.strip())[:50] or "silabo"


def _public_asset_path(filename: str) -> Path:
    current_file = Path(__file__).resolve()
    backend_root = current_file.parent.parent
    project_root = backend_root.parent

    candidates: list[Path] = []
    for env_name in (
        "SILABOS_PUBLIC_ASSETS_DIR",
        "PUBLIC_ASSETS_DIR",
        "FRONTEND_PUBLIC_DIR",
        "CPANEL_PUBLIC_HTML",
    ):
        env_value = os.getenv(env_name, "").strip()
        if env_value:
            candidates.append(Path(env_value) / filename)

    candidates.extend(
        [
            backend_root / "static" / filename,
            backend_root / "public" / filename,
            project_root / "silabos-frontend" / "public" / filename,
            project_root / "public" / filename,
            project_root / "public_html" / filename,
            Path.cwd() / filename,
            Path.cwd() / "public" / filename,
            Path.cwd() / "static" / filename,
        ]
    )

    for parent in [current_file.parent, *current_file.parents]:
        candidates.extend(
            [
                parent / filename,
                parent / "public" / filename,
                parent / "public_html" / filename,
                parent / "static" / filename,
                parent / "silabos.innovasaber.com.pe" / filename,
            ]
        )

    seen: set[str] = set()
    for candidate in candidates:
        try:
            resolved = candidate.resolve()
        except OSError:
            resolved = candidate
        key = str(resolved).lower()
        if key in seen:
            continue
        seen.add(key)
        if resolved.exists():
            return resolved

    fallback = project_root / "silabos-frontend" / "public" / filename
    logger.warning(
        "Asset publico no encontrado: %s. Configure SILABOS_PUBLIC_ASSETS_DIR con la carpeta donde estan los logos.",
        filename,
    )
    return fallback


def _clean_program_label(value: Any) -> str:
    text = _val(value)
    if not text:
        return ""
    text = re.sub(r"(?<=[a-záéíóúñ])(?=[A-ZÁÉÍÓÚÑ])", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _non_empty(value: Any) -> bool:
    return value is not None and value != "" and value != [] and value != {}


def _overlay_non_empty(base: dict[str, Any], overlay: dict[str, Any]) -> dict[str, Any]:
    merged = dict(base or {})
    for key, value in (overlay or {}).items():
        if _non_empty(value):
            merged[key] = value
    return merged


def _as_week_int(value: Any) -> int | None:
    if isinstance(value, int):
        return value
    text = _val(value)
    if not text:
        return None
    match = re.search(r"\d+", text)
    if not match:
        return None
    try:
        return int(match.group(0))
    except (TypeError, ValueError):
        return None


def _as_text_items(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        items: list[str] = []
        for item in value:
            items.extend(_as_text_items(item))
        return items
    if isinstance(value, dict):
        if isinstance(value.get("items"), list):
            return _as_text_items(value.get("items"))
        text = _val(
            value.get("name")
            or value.get("nombre")
            or value.get("title")
            or value.get("titulo")
            or value.get("descripcion")
            or value.get("statement")
            or value.get("text")
        )
        return [text] if text else []
    if isinstance(value, str):
        return [part.strip(" -\t") for part in re.split(r"[\n;]+", value) if part.strip(" -\t")]
    text = _val(value)
    return [text] if text else []


def _merge_text_items(values: list[Any], limit: int = 16) -> list[str]:
    merged: list[str] = []
    seen: set[str] = set()
    for value in values:
        for item in _as_text_items(value):
            key = re.sub(r"\s+", " ", item.lower()).strip()
            if key and key not in seen:
                seen.add(key)
                merged.append(item)
            if len(merged) >= limit:
                return merged
    return merged


def _week_range_text(values: Any) -> str:
    if isinstance(values, str):
        parts = re.findall(r"\d+", values)
        numbers = [int(part) for part in parts]
    elif isinstance(values, list):
        numbers = [
            week
            for week in (_as_week_int(value) for value in values)
            if week is not None
        ]
    else:
        week = _as_week_int(values)
        numbers = [week] if week is not None else []
    if not numbers:
        return ""
    start, end = min(numbers), max(numbers)
    return str(start) if start == end else f"{start}-{end}"


def _clean_export_unit_title(value: Any, unit_number: int) -> str:
    title = _val(value)
    title = re.sub(r"^\s*(?:unidad|u)\s*(?:[ivxlcdm]+|\d+)\s*[:.\-–—]?\s*", "", title, flags=re.IGNORECASE)
    title = re.sub(r"\s+", " ", title).strip(" .:-")
    if not title or title.lower() in {f"unidad {unit_number}", f"u {unit_number}", f"u{unit_number}"}:
        return ""
    return title


def _normalize_schedule_rows(rows: Any) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    if not isinstance(rows, list):
        return normalized
    for row in rows:
        if not isinstance(row, dict):
            continue
        week_number = _as_week_int(row.get("semana") or row.get("week"))
        if week_number is None:
            continue
        unit_number = _as_week_int(row.get("unit_number") or row.get("unidad") or row.get("unit"))
        skill_items = _merge_text_items(
            [
                row.get("habilidades_requeridas"),
                row.get("required_skills"),
                row.get("skill"),
                row.get("habilidad"),
            ],
            limit=8,
        )
        normalized_row = dict(row)
        normalized_row.update(
            {
                "semana": week_number,
                "week": week_number,
                "unit_number": unit_number,
                "desempeno": _val(row.get("desempeno") or row.get("performance")),
                "habilidades_requeridas": skill_items,
                "conocimientos": row.get("conocimientos") or row.get("knowledge") or row.get("tema") or "",
                "tema": row.get("tema") or row.get("knowledge") or row.get("conocimientos") or "",
                "actividad": row.get("actividad") or row.get("activity") or "",
                "evidencia": row.get("evidencia") or row.get("evidence") or row.get("producto") or "",
                "producto": row.get("producto") or row.get("evidence") or row.get("evidencia") or "",
                "fecha": _val(row.get("fecha") or row.get("date_range") or row.get("date")),
            }
        )
        normalized.append(normalized_row)
    return sorted(normalized, key=lambda item: int(item.get("semana") or 0))


def _apply_week_dates_to_schedule(payload: dict[str, Any], schedule_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not schedule_rows:
        return schedule_rows

    dg = payload.get("datos_generales") if isinstance(payload.get("datos_generales"), dict) else {}
    semester = (
        dg.get("semestre")
        or dg.get("periodo_academico")
        or payload.get("semester")
        or payload.get("semestre")
    )
    start_date = dg.get("fecha_inicio") or payload.get("fecha_inicio")
    max_week = max([SEMESTER_WEEKS, *[int(row.get("semana") or 0) for row in schedule_rows]])
    week_dates = _compute_week_dates(semester, total_weeks=max_week, start_date=start_date)

    dated_rows: list[dict[str, Any]] = []
    for row in schedule_rows:
        week = int(row.get("semana") or row.get("week") or 0)
        next_row = dict(row)
        next_row["fecha"] = _week_date_from_row(next_row, week, week_dates)
        next_row["date_range"] = next_row["fecha"]
        dated_rows.append(next_row)

    if isinstance(dg, dict) and start_date and not _val(dg.get("fecha_fin")):
        dg = dict(dg)
        dg["fecha_fin"] = _compute_end_date(start_date, total_weeks=SEMESTER_WEEKS)
        payload["datos_generales"] = dg

    return dated_rows


def _normalize_units(units: Any, schedule_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    unit_rows: dict[int, list[dict[str, Any]]] = {}
    for row in schedule_rows:
        unit_number = _as_week_int(row.get("unit_number"))
        if unit_number is not None:
            unit_rows.setdefault(unit_number, []).append(row)

    source_units = units if isinstance(units, list) else []
    if not source_units and unit_rows:
        source_units = [{"unit_number": unit_number} for unit_number in sorted(unit_rows)]

    normalized: list[dict[str, Any]] = []
    for index, unit in enumerate(source_units):
        if not isinstance(unit, dict):
            continue
        unit_number = _as_week_int(unit.get("unit_number") or unit.get("numero") or unit.get("unidad")) or (index + 1)
        rows_for_unit = unit_rows.get(unit_number, [])
        week_values = unit.get("weeks") or unit.get("semanas")
        if not week_values and rows_for_unit:
            week_values = [row.get("semana") for row in rows_for_unit]
        exact_week_values = unit.get("weeks") if isinstance(unit.get("weeks"), list) else []
        if not exact_week_values and rows_for_unit:
            exact_week_values = [row.get("semana") for row in rows_for_unit]
        week_numbers = [
            week
            for week in (_as_week_int(value) for value in exact_week_values)
            if week is not None
        ]
        semanas = _val(unit.get("semanas")) or _week_range_text(week_values)
        performance = _val(
            unit.get("logro")
            or unit.get("ra_unidad")
            or unit.get("desempeno")
            or unit.get("performance")
            or (rows_for_unit[0].get("desempeno") if rows_for_unit else "")
        )
        skills = unit.get("habilidades_requeridas") or unit.get("required_skills")
        if not _as_text_items(skills):
            skills = _merge_text_items([row.get("habilidades_requeridas") for row in rows_for_unit], limit=12)
        temas = unit.get("temas")
        if not temas and rows_for_unit:
            temas = [
                {
                    "semana": row.get("semana"),
                    "conocimientos": row.get("conocimientos") or row.get("tema") or "",
                }
                for row in rows_for_unit
            ]
        normalized_unit = dict(unit)
        normalized_unit.update(
            {
                "unit_number": unit_number,
                "numero": _val(unit.get("numero"), str(unit_number)),
                "titulo": _clean_export_unit_title(unit.get("titulo"), unit_number),
                "weeks": unit.get("weeks") if isinstance(unit.get("weeks"), list) else week_numbers,
                "semanas": semanas,
                "logro": performance,
                "desempeno": performance,
                "ra_unidad": _val(unit.get("ra_unidad"), performance),
                "required_skills": skills,
                "habilidades_requeridas": skills,
                "temas": temas or [],
            }
        )
        normalized.append(normalized_unit)
    return normalized


def _percentage_value(value: Any) -> Any:
    text = _val(value)
    if not text:
        return 0
    match = re.search(r"\d+(?:[.,]\d+)?", text)
    if not match:
        return value
    number_text = match.group(0).replace(",", ".")
    try:
        number = float(number_text)
    except ValueError:
        return value
    return int(number) if number.is_integer() else number


def _criteria_from_grading_payload(grading_payload: Any) -> list[dict[str, Any]]:
    rows: Any = []
    if isinstance(grading_payload, dict):
        rows = grading_payload.get("rows") or grading_payload.get("criterios") or []
    elif isinstance(grading_payload, list):
        rows = grading_payload
    if not isinstance(rows, list):
        return []
    criteria: list[dict[str, Any]] = []
    for index, row in enumerate(rows):
        if not isinstance(row, dict):
            continue
        week = _as_week_int(row.get("week") or row.get("semana") or row.get("cronograma"))
        criteria.append(
            {
                "nombre": _val(
                    row.get("nombre")
                    or row.get("name")
                    or row.get("evidencia")
                    or row.get("evidence"),
                    f"Evaluacion {index + 1}",
                ),
                "sigla": _val(row.get("sigla") or row.get("code") or row.get("label"), f"EV{index + 1}"),
                "porcentaje": _percentage_value(row.get("porcentaje") or row.get("percentage") or row.get("peso")),
                "cronograma": _val(row.get("cronograma") or row.get("schedule"), f"Semana {week}" if week else "-"),
            }
        )
    return criteria


def _prepare_export_payload(silabo: dict) -> dict:
    payload = dict(silabo or {})
    final_syllabus = payload.get("final_syllabus")
    if isinstance(final_syllabus, dict):
        payload = _overlay_non_empty(payload, final_syllabus)

    progressive = payload.get("progressive_curriculum")
    if isinstance(progressive, dict):
        prefer_progressive = isinstance(final_syllabus, dict)
        if (prefer_progressive or not payload.get("cronograma_semanal")) and isinstance(progressive.get("content_plan"), list):
            payload["cronograma_semanal"] = progressive.get("content_plan")
        if (prefer_progressive or not payload.get("unidades_tematicas")) and isinstance(progressive.get("units"), list):
            payload["unidades_tematicas"] = progressive.get("units")

    schedule_rows = _normalize_schedule_rows(payload.get("cronograma_semanal") or [])
    schedule_rows = _apply_week_dates_to_schedule(payload, schedule_rows)
    if schedule_rows:
        payload["cronograma_semanal"] = schedule_rows
    payload["unidades_tematicas"] = _normalize_units(payload.get("unidades_tematicas") or [], schedule_rows)

    sistema = payload.get("sistema_evaluacion")
    if not isinstance(sistema, dict):
        sistema = {}
    if not sistema.get("criterios"):
        criteria = _criteria_from_grading_payload(payload.get("grading"))
        if criteria:
            sistema = dict(sistema)
            sistema["criterios"] = criteria
            payload["sistema_evaluacion"] = sistema
    return payload


def _build_context(silabo: dict) -> dict:
    """
    Construye el contexto de exportación del sílabo.
    Con rowspan en exportación: desempeño y habilidades se muestran una vez por unidad.
    Matriz programa de contenidos: 6 cols (Desempeños|Habilidades|Semana/Fecha|K|Actividades|Evidencias).
    """
    silabo = _prepare_export_payload(silabo)
    dg = silabo.get("datos_generales") or {}
    semester_for_dates = (
        dg.get("semestre")
        or dg.get("periodo_academico")
        or silabo.get("semester")
        or silabo.get("semestre")
    )
    week_dates = _compute_week_dates(
        semester_for_dates,
        total_weeks=SEMESTER_WEEKS,
        start_date=dg.get("fecha_inicio") or silabo.get("fecha_inicio"),
    )
    fecha_fin = _val(dg.get("fecha_fin")) or _compute_end_date(
        dg.get("fecha_inicio") or silabo.get("fecha_inicio"),
        total_weeks=SEMESTER_WEEKS,
    )

    def _list_to_lines(value: Any) -> str:
        if isinstance(value, list):
            parts: list[str] = []
            for item in value:
                if isinstance(item, dict):
                    if isinstance(item.get("items"), list):
                        parts.extend(str(v).strip() for v in item["items"] if str(v or "").strip())
                    else:
                        text = _val(
                            item.get("name")
                            or item.get("nombre")
                            or item.get("descripcion")
                            or item.get("statement")
                        )
                        if text:
                            parts.append(text)
                else:
                    text = str(item or "").strip()
                    if text:
                        parts.append(text)
            return "\n".join(parts) if parts else "-"
        if isinstance(value, str):
            return value.strip() or "-"
        return "-"

    unidades_raw = silabo.get("unidades_tematicas") or []
    desempenos = []
    for i, unidad in enumerate(unidades_raw):
        logro = _val(unidad.get("logro"))
        desempenos.append(
            {
                "codigo": f"D{i + 1}",
                "descripcion": logro or f"D{i + 1}",
            }
        )

    ra_unidades = []
    for i, unidad in enumerate(unidades_raw):
        ra_text = _val(unidad.get("ra_unidad")) or _val(unidad.get("logro"))
        ra_unidades.append(
            {
                "codigo": f"RA{i + 1}",
                "descripcion": ra_text or f"RA{i + 1}",
            }
        )

    cronograma_map: dict[int, dict] = {}
    for item in silabo.get("cronograma_semanal") or []:
        if isinstance(item, dict) and "semana" in item:
            try:
                cronograma_map[int(item["semana"])] = item
            except (TypeError, ValueError):
                continue

    unidades_ctx = []
    roman_units = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII"]
    for i, unidad in enumerate(unidades_raw):
        desempeno_txt = _val(unidad.get("logro"), f"D{i + 1}")
        habilidades_raw = unidad.get("required_skills") or unidad.get("habilidades_requeridas")
        habilidades_txt = _list_to_lines(habilidades_raw) if habilidades_raw else _val(unidad.get("logro"), "-")

        semanas: list[int] = []
        if isinstance(unidad.get("weeks"), list):
            semanas = [
                week
                for week in (_as_week_int(value) for value in unidad.get("weeks", []))
                if week is not None
            ]
        semanas_str = _val(unidad.get("semanas"), "")
        if not semanas:
            partes = re.findall(r"\d+", semanas_str)
            if len(partes) >= 2:
                semanas = list(range(int(partes[0]), int(partes[1]) + 1))
            elif len(partes) == 1:
                semanas = [int(partes[0])]

        filas = []
        if semanas:
            for semana in semanas:
                sem_data = cronograma_map.get(semana) or {}
                temas_raw = unidad.get("temas") or []
                conocimientos = ""

                if sem_data.get("tema"):
                    conocimientos = _val(sem_data.get("tema"))
                elif temas_raw:
                    for tema in temas_raw:
                        if (
                            isinstance(tema, dict)
                            and str(tema.get("semana", "")).isdigit()
                            and int(tema["semana"]) == semana
                        ):
                            conocimientos = _val(
                                tema.get("conocimientos") or tema.get("tema")
                            )
                            break
                    if not conocimientos and isinstance(temas_raw[0], str):
                        conocimientos = ", ".join(str(t) for t in temas_raw)

                desempeno_celda = _val(sem_data.get("desempeno"), f"D{i + 1}. {desempeno_txt}")
                conocimientos_celda = _list_to_lines(sem_data.get("conocimientos")) if sem_data.get("conocimientos") else (conocimientos or "-")
                habilidades_celda = habilidades_txt
                evidencia_celda = _val(sem_data.get("evidencia") or sem_data.get("producto"), "-")
                filas.append(
                    {
                        "desempeno": desempeno_celda,
                        "habilidades": habilidades_celda,
                        "semana": f"{semana}",
                        "fecha": _week_date_from_row(sem_data, semana, week_dates),
                        "conocimientos": conocimientos_celda,
                        "actividades": _val(sem_data.get("actividad"), "-"),
                        "evidencias": evidencia_celda,
                    }
                )
        else:
            temas = unidad.get("temas") or []
            temas_txt = ", ".join(
                t if isinstance(t, str) else _val(t.get("conocimientos") or t.get("tema"))
                for t in temas
            )
            filas.append(
                {
                    "desempeno": f"D{i + 1}. {desempeno_txt}",
                    "habilidades": habilidades_txt,
                    "semana": semanas_str or "-",
                    "fecha": _week_date_from_row({}, semanas[0], week_dates) if semanas else "-",
                    "conocimientos": temas_txt or "-",
                    "actividades": "-",
                    "evidencias": "-",
                }
            )

        numero = _val(unidad.get("numero"), str(i + 1))
        titulo = _val(unidad.get("titulo"), f"Unidad {i + 1}")
        unidades_ctx.append(
            {
                "numero": numero,
                "titulo": titulo,
                "titulo_completo": f"6.{numero}. UNIDAD {roman_units[i] if i < len(roman_units) else numero}: {titulo}",
                "ra_unidad": _val(unidad.get("ra_unidad") or unidad.get("logro"), "-"),
                "desempeno": f"D{i + 1}. {desempeno_txt}",
                "habilidades": habilidades_txt,
                "filas": filas,
            }
        )

    matriz_raw = silabo.get("evaluacion_matriz") or []
    if matriz_raw:
        eval_filas = []
        for index, item in enumerate(matriz_raw):
            if not isinstance(item, dict):
                continue
            habilidades = [h for h in _list_to_lines(item.get("habilidades")).split("\n") if h and h != "-"] or ["-"]
            for habilidad in habilidades:
                eval_filas.append(
                    {
                        "desempenos": _val(
                            item.get("desempenos") or item.get("desempeno"),
                            "-",
                        ),
                        "habilidades": habilidad,
                        "evidencias": _val(
                            item.get("evidenciasDeAprendizaje") or item.get("evidencias"),
                            "-",
                        ),
                        "instrumentos": _val(item.get("instrumentos"), "-"),
                    }
                )
    else:
        eval_filas = []
        for unidad in unidades_ctx:
            evidencias = []
            seen = set()
            for fila in unidad.get("filas", []):
                evidencia = _val(fila.get("evidencias"), "")
                key = evidencia.lower()
                if evidencia and key not in seen:
                    seen.add(key)
                    evidencias.append(evidencia)
            habilidades = [h for h in str(unidad.get("habilidades") or "").split("\n") if h.strip()] or ["-"]
            for habilidad in habilidades:
                eval_filas.append(
                    {
                        "desempenos": unidad.get("desempeno") or "-",
                        "habilidades": habilidad,
                        "evidencias": "; ".join(evidencias) if evidencias else "-",
                        "instrumentos": "Rúbrica analítica",
                    }
                )

    criterios = (silabo.get("sistema_evaluacion") or {}).get("criterios") or []
    grading = []
    for index, criterio in enumerate(criterios):
        if not isinstance(criterio, dict):
            continue
        grading.append(
            {
                "evidencia": _val(criterio.get("nombre"), f"Evaluación {index + 1}"),
                "sigla": _val(criterio.get("sigla"), f"EV{index + 1}"),
                "peso": f"{criterio.get('porcentaje', 0)}%",
                "cronograma": _val(
                    criterio.get("cronograma"),
                    criterio.get("descripcion"),
                ),
            }
        )

    if not grading:
        grading = [
            {
                "evidencia": "Tareas",
                "sigla": "TA",
                "peso": "15%",
                "cronograma": "Permanente",
            },
            {
                "evidencia": "Producto Acreditable 1",
                "sigla": "PA1",
                "peso": "15%",
                "cronograma": "Semana 4",
            },
            {
                "evidencia": "Producto Acreditable 2",
                "sigla": "PA2",
                "peso": "20%",
                "cronograma": "Semana 8",
            },
            {
                "evidencia": "Examen Parcial",
                "sigla": "EP",
                "peso": "15%",
                "cronograma": "Semana 12",
            },
            {
                "evidencia": "Producto Acreditable 3",
                "sigla": "PA3",
                "peso": "35%",
                "cronograma": "Semana 16",
            },
        ]

    referencias = [
        _val(b.get("referencia"))
        for b in (silabo.get("bibliografia") or [])
        if isinstance(b, dict) and b.get("referencia")
    ]

    competencia = _val(
        silabo.get("competencia_profesional")
        or (silabo.get("competencias") or [""])[0]
    )
    capacidad = _val(
        silabo.get("capacidad_del_curso")
        or (silabo.get("resultados_aprendizaje") or [""])[0]
    )
    ra_curso = _val(
        silabo.get("resultado_aprendizaje")
        or (silabo.get("resultados_aprendizaje") or [""])[0]
    )
    responsabilidad_raw = (
        silabo.get("responsabilidad_social")
        or silabo.get("responsabilidadSocial")
        or {}
    )
    responsabilidad_fallback = (
        "Como actividad de responsabilidad social, los estudiantes identificarán una necesidad concreta del entorno universitario o comunitario vinculada al propósito del curso.\n"
        "Organizados en equipos, recogerán información breve del contexto y la interpretarán a partir de los desempeños oficiales.\n"
        "Con base en ese análisis, diseñarán una acción formativa, demostrativa o de orientación que aplique los aprendizajes del curso ante destinatarios reales.\n"
        "La actividad se cerrará con una evidencia verificable, evitando acciones decorativas o aisladas."
    )
    if isinstance(responsabilidad_raw, dict):
        responsabilidad_social = _val(
            responsabilidad_raw.get("actividadPropuesta")
            or responsabilidad_raw.get("actividad_propuesta")
            or responsabilidad_raw.get("descripcion"),
            responsabilidad_fallback,
        )
    else:
        responsabilidad_social = _val(
            responsabilidad_raw,
            responsabilidad_fallback,
        )

    return {
        "facultad": _clean_program_label(dg.get("facultad")),
        "escuela_profesional": _clean_program_label(
            dg.get("escuela_profesional") or dg.get("carrera")
        ),
        "programa_estudios": _clean_program_label(
            dg.get("programa_estudios") or dg.get("carrera")
        ),
        "departamento_academico": _clean_program_label(
            dg.get("escuela_profesional") or dg.get("carrera")
        ),
        "nombre_curso": _val(dg.get("nombre_curso")),
        "programa": _clean_program_label(dg.get("programa_estudios") or dg.get("carrera")),
        "escuela": _clean_program_label(dg.get("escuela_profesional") or dg.get("carrera")),
        "modalidad": _val(dg.get("modalidad"), "Presencial").capitalize(),
        "curso": _val(dg.get("nombre_curso")),
        "prerrequisito": _prerequisite_value(dg.get("prerrequisito")),
        "codigo_curso": _val(dg.get("codigo")),
        "semestre": _val(dg.get("semestre")),
        "periodo_academico": _val(
            dg.get("periodo_academico") or dg.get("semestre")
        ),
        "creditos": _val(dg.get("creditos")),
        "horas_teoria": _format_hour_value(dg.get("horas_teoria")),
        "horas_practica": _format_hour_value(dg.get("horas_practica")),
        "fecha_inicio": _val(dg.get("fecha_inicio")),
        "fecha_fin": fecha_fin,
        "horas_semanales": _format_hours(dg.get("horas_teoria"), dg.get("horas_practica")),
        "docente_nombre": _val(dg.get("docente")),
        "docente_email": _val(dg.get("docente_email")),
        "sumilla": _val(silabo.get("sumilla")),
        "competencia_profesional": competencia,
        "capacidad_del_curso": capacidad,
        "ra_curso": ra_curso,
        "ra_unidades": ra_unidades,
        "desempenos": desempenos,
        "unidades": unidades_ctx,
        "eval_filas": eval_filas,
        "grading": grading,
        "formula_pf": "PF = " + " + ".join(f"{row['peso']}*{row['sigla']}" for row in grading),
        "metodologia": _sanitize_sentence_spacing(_val(silabo.get("metodologia"), "")),
        "tutoria": _sanitize_sentence_spacing(_val(silabo.get("tutoria"), "")),
        "responsabilidad_social": _sanitize_sentence_spacing(responsabilidad_social),
        "referencias": referencias,
        "_filename": _safe_filename(_val(dg.get("nombre_curso"))),
    }


def _add_section_title(document, title: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True


def _add_blank_line(document, points: int = 8) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.add_run(" ")


def _add_paragraph_block(document, text: str) -> None:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", str(text or "")) if part.strip()]
    if not paragraphs:
        document.add_paragraph("-")
        return
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text or "—")
    run.bold = bold


def _set_table_borders(table, *, color: str = "FFFFFF", size: str = "4") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def apply_borders(container, tag_name: str) -> None:
        borders = container.find(qn(tag_name))
        if borders is None:
            borders = OxmlElement(tag_name)
            container.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge_tag = f"w:{edge}"
            node = borders.find(qn(edge_tag))
            if node is None:
                node = OxmlElement(edge_tag)
                borders.append(node)
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), size)
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), color)

    apply_borders(table._tbl.tblPr, "w:tblBorders")
    for row in table.rows:
        for cell in row.cells:
            apply_borders(cell._tc.get_or_add_tcPr(), "w:tcBorders")


def _merge_repeated_vertical_cells(table, col_idx: int, row_start: int, row_end: int, text: str) -> None:
    if row_end <= row_start:
        _set_cell_text(table.cell(row_start, col_idx), text)
        return
    merged = table.cell(row_start, col_idx).merge(table.cell(row_end, col_idx))
    _set_cell_text(merged, text)


def _generar_docx_programatico(context: dict) -> bytes:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)
    header_table = document.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    _set_table_borders(header_table)
    header_cells = header_table.rows[0].cells
    header_cells[0].width = Cm(3.4)
    header_cells[1].width = Cm(10.2)
    header_cells[2].width = Cm(3.4)
    logo_unprg = _public_asset_path("unprg-logo.png")
    logo_fachse = _public_asset_path("logo_fachse.png")
    if logo_unprg.exists():
        header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[0].paragraphs[0].add_run().add_picture(str(logo_unprg), width=Cm(1.65))
    header = header_cells[1].paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('UNIVERSIDAD NACIONAL "PEDRO RUIZ GALLO"\n')
    run.bold = True
    run.font.size = Pt(13)
    faculty_name = context["facultad"] or "................................"
    faculty_label = faculty_name if faculty_name.lower().startswith("facultad") else f"Facultad de {faculty_name}"
    run = header.add_run(f"{faculty_label}\n")
    run.bold = True
    run = header.add_run(f"ESCUELA PROFESIONAL {context['escuela_profesional'] or '................................'}\n")
    run.bold = True
    header.add_run(f"Departamento Académico de {context['departamento_academico'] or '................................'}")
    if logo_fachse.exists():
        header_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[2].paragraphs[0].add_run().add_picture(str(logo_fachse), width=Cm(1.65))
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{context['nombre_curso']} (Sílabo)")
    run.bold = True

    _add_section_title(document, "I. Información General")
    info_table = document.add_table(rows=0, cols=3)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    _set_table_borders(info_table)
    info_fields = [
        ("1.1.   Programa de Estudios", context["programa"]),
        ("1.2.   Escuela Profesional", context["escuela"]),
        ("1.3.   Modalidad", context["modalidad"]),
        ("1.4.   Curso", context["curso"]),
        ("1.5.   Prerrequisito", context["prerrequisito"]),
        ("1.6.   Código del curso", context["codigo_curso"]),
        ("1.7.   Semestre Académico", context["semestre"]),
        ("1.8.   Periodo Académico", context["periodo_academico"]),
        ("1.9.   Créditos", context["creditos"]),
        ("1.10.  Horas Semanales", ""),
        ("        Teoría", context["horas_teoria"]),
        ("        Práctica", context["horas_practica"]),
        ("1.11.  Duración", ""),
        ("        Fecha de inicio", context["fecha_inicio"]),
        ("        Fecha de término", context["fecha_fin"]),
        (
            "1.12.  Docente",
            _pair(context["docente_nombre"], context["docente_email"], "").replace(" / ", "\n"),
        ),
    ]
    for label, value in info_fields:
        row = info_table.add_row().cells
        row[0].width = Cm(8)
        row[1].width = Cm(0.35)
        row[2].width = Cm(8)
        _set_cell_text(row[0], label)
        _set_cell_text(row[1], ":")
        _set_cell_text(row[2], value or " ")
    _set_table_borders(info_table)

    _add_section_title(document, "II. Sumilla")
    document.add_paragraph(context["sumilla"] or "—")

    _add_section_title(document, "III. Competencia Profesional")
    document.add_paragraph(context["competencia_profesional"] or "—")

    _add_section_title(document, "IV. Capacidad del Curso")
    document.add_paragraph(context["capacidad_del_curso"] or context["ra_curso"] or "—")

    _add_section_title(document, "V. Desempeños de las Unidades Didácticas")
    if context["desempenos"]:
        for desempeno in context["desempenos"]:
            document.add_paragraph(
                f"{desempeno['codigo']}. {desempeno['descripcion']}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("—")

    _add_blank_line(document, 10)
    _add_section_title(document, "VI. Programa de Contenidos")
    for unidad in context["unidades"]:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(unidad["titulo_completo"])
        run.bold = True
        unit_table = document.add_table(rows=1, cols=6)
        unit_table.style = "Table Grid"
        unit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = [
            "Desempeños",
            "Habilidades requeridas",
            "Semana (Fecha)",
            "Conocimientos",
            "Actividades",
            "Evidencias de Aprendizaje",
        ]
        for index, header_text in enumerate(headers):
            _set_cell_text(unit_table.rows[0].cells[index], header_text, bold=True)

        for fila in unidad["filas"]:
            row = unit_table.add_row().cells
            _set_cell_text(row[0], fila["desempeno"])
            _set_cell_text(row[1], fila["habilidades"])
            _set_cell_text(row[2], f"Sem. {fila['semana']}\n{fila.get('fecha', '—')}")
            _set_cell_text(row[3], fila["conocimientos"])
            _set_cell_text(row[4], fila["actividades"])
            _set_cell_text(row[5], fila["evidencias"])

        last_row = len(unidad.get("filas", []))
        if last_row >= 1:
            _merge_repeated_vertical_cells(unit_table, 0, 1, last_row, unidad.get("desempeno", "—"))
            _merge_repeated_vertical_cells(unit_table, 1, 1, last_row, unidad.get("habilidades", "—"))

    _add_blank_line(document, 10)
    _add_section_title(document, "VII. Sistema de Evaluación")
    eval_table = document.add_table(rows=1, cols=4)
    eval_table.style = "Table Grid"
    eval_headers = [
        "Desempeños",
        "Habilidades requeridas",
        "Evidencias de Aprendizaje",
        "Instrumentos",
    ]
    for index, header_text in enumerate(eval_headers):
        _set_cell_text(eval_table.rows[0].cells[index], header_text, bold=True)
    for fila in context["eval_filas"]:
        row = eval_table.add_row().cells
        _set_cell_text(row[0], fila["desempenos"])
        _set_cell_text(row[1], fila.get("habilidades", "—"))
        _set_cell_text(row[2], fila["evidencias"])
        _set_cell_text(row[3], fila["instrumentos"])

    current = 1
    while current < len(eval_table.rows):
        value = eval_table.cell(current, 0).text
        end = current
        while end + 1 < len(eval_table.rows) and eval_table.cell(end + 1, 0).text == value:
            end += 1
        if end > current:
            _merge_repeated_vertical_cells(eval_table, 0, current, end, value)
        current = end + 1

    _add_section_title(document, "VIII. Sistema de Calificación")
    grading_table = document.add_table(rows=1, cols=4)
    grading_table.style = "Table Grid"
    for index, header_text in enumerate(
        ["Evidencias", "Sigla", "Peso", "Cronograma"]
    ):
        _set_cell_text(grading_table.rows[0].cells[index], header_text, bold=True)
    for row_data in context["grading"]:
        row = grading_table.add_row().cells
        _set_cell_text(row[0], row_data["evidencia"])
        _set_cell_text(row[1], row_data["sigla"])
        _set_cell_text(row[2], row_data["peso"])
        _set_cell_text(row[3], row_data["cronograma"])
    paragraph = document.add_paragraph()
    run = paragraph.add_run(context["formula_pf"])
    run.bold = True

    _add_section_title(
        document,
        "IX. Metodología de Enseñanza-Aprendizaje y Actividades de Investigación Formativa",
    )
    document.add_paragraph(context["metodologia"] or "—")

    _add_section_title(document, "X. Actividades de Tutoría: Área Académica")
    document.add_paragraph(context["tutoria"] or "—")

    _add_section_title(document, "XI. Responsabilidad Social")
    document.add_paragraph(context["responsabilidad_social"] or "—")

    _add_section_title(document, "XII. Referencias")
    if context["referencias"]:
        for referencia in context["referencias"]:
            document.add_paragraph(referencia, style="List Bullet")
    else:
        document.add_paragraph("Pendiente de completar.")

    signature_table = document.add_table(rows=2, cols=2)
    signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(signature_table.rows[0].cells[0], "__________________________")
    _set_cell_text(signature_table.rows[0].cells[1], "__________________________")
    _set_cell_text(
        signature_table.rows[1].cells[0],
        "Director del Departamento Académico",
    )
    _set_cell_text(
        signature_table.rows[1].cells[1],
        context["docente_nombre"] or "Docente",
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generar_docx(silabo: dict) -> bytes:
    """Genera el DOCX del silabo con el constructor programatico estable."""
    context = _build_context(silabo)
    return _generar_docx_programatico(context)


def _ensure_weasyprint_compatibility() -> None:
    import pydyf
    from weasyprint.matrix import Matrix
    from weasyprint.pdf.stream import Stream

    if not hasattr(pydyf.Stream, "transform") and not getattr(
        Stream, "_silabos_transform_patch", False
    ):

        def _patched_transform(self, a=1, b=0, c=0, d=1, e=0, f=0):
            pydyf.Stream.set_matrix(self, a, b, c, d, e, f)
            self._ctm_stack[-1] = Matrix(a, b, c, d, e, f) @ self.ctm

        Stream.transform = _patched_transform
        Stream._silabos_transform_patch = True

    if not hasattr(pydyf.Stream, "text_matrix") and not getattr(
        Stream, "_silabos_text_matrix_patch", False
    ):

        def _patched_text_matrix(self, a=1, b=0, c=0, d=1, e=0, f=0):
            pydyf.Stream.set_text_matrix(self, a, b, c, d, e, f)

        Stream.text_matrix = _patched_text_matrix
        Stream._silabos_text_matrix_patch = True


def generar_pdf_html(silabo: dict) -> bytes:
    """Genera el PDF del sílabo usando WeasyPrint desde HTML."""
    try:
        from weasyprint import CSS, HTML
    except ImportError as exc:
        raise ImportError(
            "WeasyPrint no está instalado. Ver instrucciones en requirements.txt."
        ) from exc

    _ensure_weasyprint_compatibility()

    context = _build_context(silabo)
    html_str = _build_html(context)
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    return HTML(string=html_str, base_url=base_dir).write_pdf(
        stylesheets=[CSS(string=_CSS_SILABO)]
    )


_CSS_SILABO = """
@page {
    size: A4;
    margin: 2cm 2.5cm 2cm 2.5cm;
}
body {
    font-family: Arial, sans-serif;
    font-size: 10pt;
    line-height: 1.5;
    color: #1a1a1a;
}
h1 { font-size: 13pt; text-align: center; text-transform: uppercase; margin: 0; }
h2 { font-size: 11pt; text-align: center; margin: 2px 0; }
h3 { font-size: 10pt; text-align: center; margin: 2px 0; }
p { margin: 4px 0; }
ul { margin: 4px 0; padding-left: 16px; }
li { margin: 2px 0; }
.header { text-align: center; margin-bottom: 20px; }
.header-logos {
    display: flex;
    justify-content: space-between;
    align-items: center;
    margin-bottom: 10px;
}
.header-logos img { height: 60px; width: auto; }
.header-text { flex: 1; text-align: center; }
.titulo-silabo {
    background: #f0f0f0;
    padding: 8px;
    text-align: center;
    font-weight: bold;
    font-size: 12pt;
    margin: 15px 0;
    text-transform: uppercase;
}
.seccion-titulo {
    font-weight: bold;
    font-size: 10pt;
    text-transform: uppercase;
    border-bottom: 1px solid #333;
    padding-bottom: 3px;
    margin: 15px 0 8px 0;
}
.info-grid {
    display: table;
    width: 100%;
    font-size: 9.5pt;
}
.info-row { display: table-row; }
.info-label {
    display: table-cell;
    width: 45%;
    padding: 2px 4px;
}
.info-value {
    display: table-cell;
    padding: 2px 4px;
}
table {
    width: 100%;
    border-collapse: collapse;
    font-size: 8.5pt;
    margin-bottom: 10px;
}
th {
    background: #e8e8e8;
    border: 1px solid #333;
    padding: 4px 6px;
    text-align: left;
    font-weight: bold;
}
td {
    border: 1px solid #555;
    padding: 4px 6px;
    vertical-align: top;
}
.unidad-titulo {
    font-weight: bold;
    margin: 10px 0 4px 0;
    font-size: 9.5pt;
}
.formula {
    font-weight: bold;
    margin-top: 6px;
    font-size: 9.5pt;
}
.firmas {
    margin-top: 40px;
    display: flex;
    justify-content: space-between;
}
.firma-bloque { text-align: center; width: 45%; }
.firma-linea {
    border-top: 1px solid #333;
    padding-top: 4px;
    margin-top: 30px;
}
.nota-cursiva { font-style: italic; font-size: 9pt; margin-top: 4px; }
"""


def _img_tag(path: Path, alt: str) -> str:
    if not path.exists():
        return ""
    return f'<img src="{path.as_uri()}" alt="{escape(alt)}" />'


def _build_html(ctx: dict) -> str:
    """Construye el HTML del sílabo para WeasyPrint."""
    logo_unprg = _img_tag(_public_asset_path("unprg-logo.png"), "Logo UNPRG")
    logo_fachse = _img_tag(_public_asset_path("logo_fachse.png"), "Logo FACHSE")

    def _multiline_html(value: str) -> str:
        return escape(value or "—").replace("\n", "<br/>")

    tablas_vi = []
    for unidad in ctx.get("unidades", []):
        filas = []
        unit_rows = unidad.get("filas", [])
        row_span = max(len(unit_rows), 1)
        for index, fila in enumerate(unit_rows):
            merged_prefix = ""
            if index == 0:
                merged_prefix = (
                    f"<td rowspan=\"{row_span}\">{_multiline_html(unidad.get('desempeno', fila['desempeno']))}</td>"
                    f"<td rowspan=\"{row_span}\">{_multiline_html(unidad.get('habilidades', fila['habilidades']))}</td>"
                )
            filas.append(
                f"""
                <tr>
                  {merged_prefix}
                  <td>{escape(f"Sem. {fila['semana']}")}<br/>{escape(fila.get('fecha', '—'))}</td>
                  <td>{_multiline_html(fila['conocimientos'])}</td>
                  <td>{_multiline_html(fila['actividades'])}</td>
                  <td>{_multiline_html(fila['evidencias'])}</td>
                </tr>
                """
            )

        tablas_vi.append(
            f"""
            <div class="unidad-titulo">{escape(unidad.get('titulo_completo', ''))}</div>
            <table>
              <thead>
                <tr>
                  <th>Desempeños</th>
                  <th>Habilidades requeridas</th>
                  <th>Semana (Fecha)</th>
                  <th>Conocimientos</th>
                  <th>Actividades</th>
                  <th>Evidencias de Aprendizaje</th>
                </tr>
              </thead>
              <tbody>
                {''.join(filas)}
              </tbody>
            </table>
            """
        )

    eval_rows_parts = []
    eval_items = ctx.get("eval_filas", [])
    index = 0
    while index < len(eval_items):
        current = eval_items[index]
        end = index
        while end + 1 < len(eval_items) and eval_items[end + 1].get("desempenos") == current.get("desempenos"):
            end += 1
        rowspan = end - index + 1
        for row_index in range(index, end + 1):
            fila = eval_items[row_index]
            merged_prefix = ""
            if row_index == index:
                merged_prefix = (
                    f"<td rowspan=\"{rowspan}\">{escape(current['desempenos'])}</td>"
                )
            eval_rows_parts.append(
                f"""
                <tr>
                  {merged_prefix}
                  <td>{escape(fila.get('habilidades', '—'))}</td>
                  <td>{escape(fila.get('evidencias', '-'))}</td>
                  <td>{escape(fila.get('instrumentos', '-'))}</td>
                </tr>
                """
            )
        index = end + 1
    eval_rows = "".join(eval_rows_parts)

    grading_rows = "".join(
        f"""
        <tr>
          <td>{escape(row['evidencia'])}</td>
          <td>{escape(row['sigla'])}</td>
          <td>{escape(row['peso'])}</td>
          <td>{escape(row['cronograma'])}</td>
        </tr>
        """
        for row in ctx.get("grading", [])
    )

    desempenos_html = "".join(
        f"<li>{escape(d['codigo'])}. {escape(d['descripcion'])}</li>"
        for d in ctx.get("desempenos", [])
    )

    if ctx.get("referencias"):
        refs_html = "".join(
            f"<li>{escape(ref)}</li>" for ref in ctx.get("referencias", [])
        )
    else:
        refs_html = "<li>Pendiente de completar.</li>"

    return f"""
<!DOCTYPE html>
<html lang="es">
  <head>
    <meta charset="utf-8" />
    <title>Sílabo</title>
  </head>
  <body>
    <div class="header">
      <div class="header-logos">
        <div>{logo_unprg}</div>
        <div class="header-text">
          <h1>UNIVERSIDAD NACIONAL "PEDRO RUIZ GALLO"</h1>
          <h2>{escape(ctx['facultad'])}</h2>
          <h3>ESCUELA PROFESIONAL {escape(ctx['escuela_profesional'])}</h3>
          <p>Departamento Académico de {escape(ctx['departamento_academico'])}</p>
        </div>
        <div>{logo_fachse}</div>
      </div>
      <div class="titulo-silabo">{escape(ctx['nombre_curso'])} (Sílabo)</div>
    </div>

    <div class="seccion-titulo">I. Información General</div>
    <div class="info-grid">
      <div class="info-row"><div class="info-label">1.1 Programa de Estudios</div><div class="info-value">{escape(ctx['programa'])}</div></div>
      <div class="info-row"><div class="info-label">1.2 Escuela Profesional</div><div class="info-value">{escape(ctx['escuela'])}</div></div>
      <div class="info-row"><div class="info-label">1.3 Modalidad</div><div class="info-value">{escape(ctx['modalidad'])}</div></div>
      <div class="info-row"><div class="info-label">1.4 Curso</div><div class="info-value">{escape(ctx['curso'])}</div></div>
      <div class="info-row"><div class="info-label">1.5 Prerrequisito</div><div class="info-value">{escape(ctx['prerrequisito'])}</div></div>
      <div class="info-row"><div class="info-label">1.6 Código del curso</div><div class="info-value">{escape(ctx['codigo_curso'])}</div></div>
      <div class="info-row"><div class="info-label">1.7 Semestre Académico</div><div class="info-value">{escape(ctx['semestre'])}</div></div>
      <div class="info-row"><div class="info-label">1.8 Periodo Académico</div><div class="info-value">{escape(ctx['periodo_academico'])}</div></div>
      <div class="info-row"><div class="info-label">1.9 Créditos</div><div class="info-value">{escape(ctx['creditos'])}</div></div>
      <div class="info-row"><div class="info-label">1.10<br/>Horas Semanales</div><div class="info-value">{escape(ctx['horas_semanales'])}</div></div>
      <div class="info-row"><div class="info-label">1.11 Duración (Inicio / Término)</div><div class="info-value">{escape(_pair(ctx['fecha_inicio'], ctx['fecha_fin']))}</div></div>
      <div class="info-row"><div class="info-label">1.12 Docente (Nombre / Correo)</div><div class="info-value">{escape(_pair(ctx['docente_nombre'], ctx['docente_email']))}</div></div>
    </div>

    <div class="seccion-titulo">II. Sumilla</div>
    <p>{escape(ctx['sumilla'])}</p>

    <div class="seccion-titulo">III. Competencia Profesional</div>
    <p>{escape(ctx.get('competencia_profesional') or '—')}</p>

    <div class="seccion-titulo">IV. Capacidad del Curso</div>
    <p>{escape(ctx.get('capacidad_del_curso') or ctx.get('ra_curso') or '—')}</p>

    <div class="seccion-titulo">V. Desempeños de las Unidades Didácticas</div>
    <ul>{desempenos_html}</ul>

    <div class="seccion-titulo">VI. Programa de Contenidos</div>
    {''.join(tablas_vi)}

    <div class="seccion-titulo">VII. Sistema de Evaluación</div>
    <table>
      <thead>
        <tr>
          <th>Desempeños</th>
          <th>Habilidades requeridas</th>
          <th>Evidencias de Aprendizaje</th>
          <th>Instrumentos</th>
        </tr>
      </thead>
      <tbody>{eval_rows}</tbody>
    </table>

    <div class="seccion-titulo">VIII. Sistema de Calificación</div>
    <table>
      <thead>
        <tr>
          <th>Evidencias de aprendizaje</th>
          <th>Sigla</th>
          <th>Peso</th>
          <th>Cronograma</th>
        </tr>
      </thead>
      <tbody>{grading_rows}</tbody>
    </table>
    <div class="formula">{escape(ctx['formula_pf'])}</div>

    <div class="seccion-titulo">IX. Metodología de Enseñanza-Aprendizaje y Actividades de Investigación Formativa</div>
    <p>{escape(ctx['metodologia'])}</p>

    <div class="seccion-titulo">X. Actividades de Tutoría: Área Académica</div>
    <p>{escape(ctx['tutoria'])}</p>

    <div class="seccion-titulo">XI. Responsabilidad Social</div>
    <p>{escape(ctx['responsabilidad_social'])}</p>

    <div class="seccion-titulo">XII. Referencias</div>
    <ul>{refs_html}</ul>

    <div class="firmas">
      <div class="firma-bloque">
        <div class="firma-linea">Nombre y apellidos</div>
        <div>Director del Departamento Académico</div>
      </div>
      <div class="firma-bloque">
        <div class="firma-linea">{escape(ctx['docente_nombre'])}</div>
        <div>Docente</div>
      </div>
    </div>
  </body>
</html>
"""
